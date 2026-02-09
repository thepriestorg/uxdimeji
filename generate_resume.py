from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '333333')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_resume():
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # ============ HEADER ============
    # Name
    name = doc.add_paragraph()
    name_run = name.add_run("OLADIMEJI ABUBAKAR")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.name = 'Arial'
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.space_after = Pt(0)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Product Designer | UI/UX Engineer")
    title_run.font.size = Pt(12)
    title_run.font.name = 'Arial'
    title_run.font.color.rgb = RGBColor(80, 80, 80)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(6)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact_run = contact.add_run("Lagos, Nigeria | [Phone] | [email@email.com] | linkedin.com/in/yourprofile | uxdimeji.com")
    contact_run.font.size = Pt(10)
    contact_run.font.name = 'Arial'
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(6)
    add_horizontal_line(contact)
    
    # ============ PROFESSIONAL SUMMARY ============
    section_heading = doc.add_paragraph()
    section_run = section_heading.add_run("PROFESSIONAL SUMMARY")
    section_run.bold = True
    section_run.font.size = Pt(12)
    section_run.font.name = 'Arial'
    section_heading.space_before = Pt(12)
    section_heading.space_after = Pt(6)
    
    summary = doc.add_paragraph()
    summary_text = "Results-driven Product Designer with 5+ years of experience crafting intuitive, scalable digital products across consumer tech, SaaS, logistics, and GovTech sectors. First Class Computer Science graduate combining deep technical expertise with human-centered design methodology. Proven track record of translating complex requirements into elegant interfaces that drive measurable business outcomes. Skilled in cross-functional collaboration, design systems, and bridging the gap between design and development."
    summary_run = summary.add_run(summary_text)
    summary_run.font.size = Pt(10)
    summary_run.font.name = 'Arial'
    summary.space_after = Pt(6)
    add_horizontal_line(summary)
    
    # ============ CORE SKILLS ============
    section_heading = doc.add_paragraph()
    section_run = section_heading.add_run("CORE SKILLS")
    section_run.bold = True
    section_run.font.size = Pt(12)
    section_run.font.name = 'Arial'
    section_heading.space_before = Pt(12)
    section_heading.space_after = Pt(6)
    
    # Design Skills
    skills_para = doc.add_paragraph()
    label = skills_para.add_run("Design: ")
    label.bold = True
    label.font.size = Pt(10)
    label.font.name = 'Arial'
    design_skills = "User Interface (UI) Design, User Experience (UX) Design, Interaction Design, Visual Design, Design Systems & Component Libraries, Responsive & Mobile-First Design, Wireframing & Prototyping, User Research & Usability Testing, Information Architecture, Accessibility (WCAG)"
    skills_run = skills_para.add_run(design_skills)
    skills_run.font.size = Pt(10)
    skills_run.font.name = 'Arial'
    skills_para.space_after = Pt(4)
    
    # Tools
    tools_para = doc.add_paragraph()
    label = tools_para.add_run("Tools: ")
    label.bold = True
    label.font.size = Pt(10)
    label.font.name = 'Arial'
    tools = "Figma, Adobe Creative Suite (Photoshop, Illustrator, XD), Sketch, InVision, Framer, Miro, FigJam"
    tools_run = tools_para.add_run(tools)
    tools_run.font.size = Pt(10)
    tools_run.font.name = 'Arial'
    tools_para.space_after = Pt(4)
    
    # Technical Skills
    tech_para = doc.add_paragraph()
    label = tech_para.add_run("Technical: ")
    label.bold = True
    label.font.size = Pt(10)
    label.font.name = 'Arial'
    tech_skills = "HTML/CSS/JavaScript, React.js, Next.js, TypeScript, Tailwind CSS, Framer Motion, Supabase, Git & Version Control, REST APIs, No-Code/Low-Code Platforms"
    tech_run = tech_para.add_run(tech_skills)
    tech_run.font.size = Pt(10)
    tech_run.font.name = 'Arial'
    tech_para.space_after = Pt(4)
    
    # Soft Skills
    soft_para = doc.add_paragraph()
    label = soft_para.add_run("Soft Skills: ")
    label.bold = True
    label.font.size = Pt(10)
    label.font.name = 'Arial'
    soft_skills = "Cross-Functional Collaboration, Stakeholder Management, Agile Methodologies, Design Thinking, Problem Solving, Communication & Presentation"
    soft_run = soft_para.add_run(soft_skills)
    soft_run.font.size = Pt(10)
    soft_run.font.name = 'Arial'
    soft_para.space_after = Pt(6)
    add_horizontal_line(soft_para)
    
    # ============ PROFESSIONAL EXPERIENCE ============
    section_heading = doc.add_paragraph()
    section_run = section_heading.add_run("PROFESSIONAL EXPERIENCE")
    section_run.bold = True
    section_run.font.size = Pt(12)
    section_run.font.name = 'Arial'
    section_heading.space_before = Pt(12)
    section_heading.space_after = Pt(8)
    
    # Experience entries
    experiences = [
        {
            "title": "Product Designer",
            "company": "Elevayt",
            "period": "January 2026 – Present",
            "bullets": [
                "Partner with government bodies to revolutionize education systems through digital transformation initiatives",
                "Design and develop digital infrastructure powering next-generation learning platforms",
                "Collaborate with cross-functional teams including engineers, product managers, and government stakeholders",
                "Create comprehensive design systems ensuring consistency across GovTech and EdTech products",
                "Conduct user research with educators and students to inform product decisions"
            ]
        },
        {
            "title": "Contract Product Designer",
            "company": "LoopWise",
            "period": "March 2025 – June 2025",
            "bullets": [
                "Designed the mobile experience for shippers and carriers in the freight logistics industry",
                "Streamlined complex logistics workflows from trucking to air cargo into seamless tracking and booking flows",
                "Created intuitive interfaces that reduced user friction in multi-modal shipping processes",
                "Collaborated with engineering teams to ensure design feasibility and optimal implementation"
            ]
        },
        {
            "title": "Contract Product Designer",
            "company": "RoomService",
            "period": "November 2023 – February 2024",
            "bullets": [
                "Optimized food delivery user experience, focusing on the journey from discovery to checkout",
                "Implemented design improvements that contributed to increased conversion rates",
                "Conducted competitive analysis and user testing to identify optimization opportunities",
                "Designed mobile-first interfaces ensuring seamless ordering experience across devices"
            ]
        },
        {
            "title": "Contract Product Designer",
            "company": "TheClub",
            "period": "January 2023 – August 2023",
            "bullets": [
                "Designed a specialized social streaming platform for DJs to stream live sets and connect with fans",
                "Created engaging interfaces that fostered community building and digital fan engagement",
                "Developed user flows for live streaming, social interaction, and content discovery features",
                "Collaborated with development teams to implement real-time streaming functionality"
            ]
        },
        {
            "title": "Contract Product Designer",
            "company": "Alta Inc",
            "period": "November 2022 – January 2023",
            "bullets": [
                "Worked on a no-code mobile app builder empowering creators to launch products without writing code",
                "Transformed complex development logic into intuitive visual interfaces",
                "Designed drag-and-drop components and workflow builders for non-technical users"
            ]
        },
        {
            "title": "Design Intern",
            "company": "LeftLane",
            "period": "June 2022 – December 2022",
            "bullets": [
                "Contributed to consumer product design where user experience is paramount",
                "Collaborated with senior designers on B2C digital products",
                "Gained expertise in balancing speed with exceptional user experience"
            ]
        }
    ]
    
    for exp in experiences:
        # Job Title and Company
        job_line = doc.add_paragraph()
        title_run = job_line.add_run(exp["title"])
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.name = 'Arial'
        job_line.space_after = Pt(0)
        
        company_line = doc.add_paragraph()
        company_run = company_line.add_run(f"{exp['company']} | {exp['period']}")
        company_run.font.size = Pt(10)
        company_run.font.name = 'Arial'
        company_run.italic = True
        company_line.space_after = Pt(4)
        
        # Bullet points
        for bullet in exp["bullets"]:
            bullet_para = doc.add_paragraph(style='List Bullet')
            bullet_run = bullet_para.add_run(bullet)
            bullet_run.font.size = Pt(10)
            bullet_run.font.name = 'Arial'
            bullet_para.space_after = Pt(2)
            bullet_para.paragraph_format.left_indent = Inches(0.25)
        
        # Add spacing after each job
        doc.add_paragraph().space_after = Pt(4)
    
    # Add separator
    sep = doc.add_paragraph()
    add_horizontal_line(sep)
    
    # ============ EDUCATION ============
    section_heading = doc.add_paragraph()
    section_run = section_heading.add_run("EDUCATION")
    section_run.bold = True
    section_run.font.size = Pt(12)
    section_run.font.name = 'Arial'
    section_heading.space_before = Pt(8)
    section_heading.space_after = Pt(6)
    
    degree = doc.add_paragraph()
    degree_run = degree.add_run("Bachelor of Science in Computer Science")
    degree_run.bold = True
    degree_run.font.size = Pt(11)
    degree_run.font.name = 'Arial'
    degree.space_after = Pt(0)
    
    uni = doc.add_paragraph()
    uni_run = uni.add_run("[University Name] | [Graduation Year]")
    uni_run.font.size = Pt(10)
    uni_run.font.name = 'Arial'
    uni_run.italic = True
    uni.space_after = Pt(2)
    
    honors = doc.add_paragraph(style='List Bullet')
    honors_run = honors.add_run("First Class Honors")
    honors_run.font.size = Pt(10)
    honors_run.font.name = 'Arial'
    honors.paragraph_format.left_indent = Inches(0.25)
    
    coursework = doc.add_paragraph(style='List Bullet')
    cw_run = coursework.add_run("Relevant Coursework: Human-Computer Interaction, Software Engineering, Data Structures, Algorithms, Database Systems")
    cw_run.font.size = Pt(10)
    cw_run.font.name = 'Arial'
    coursework.paragraph_format.left_indent = Inches(0.25)
    coursework.space_after = Pt(6)
    add_horizontal_line(coursework)
    
    # ============ PROJECTS & ACHIEVEMENTS ============
    section_heading = doc.add_paragraph()
    section_run = section_heading.add_run("PROJECTS & ACHIEVEMENTS")
    section_run.bold = True
    section_run.font.size = Pt(12)
    section_run.font.name = 'Arial'
    section_heading.space_before = Pt(8)
    section_heading.space_after = Pt(6)
    
    projects = [
        "Portfolio Website: Designed and developed a high-end personal portfolio using Next.js, TypeScript, Supabase, and Framer Motion with custom CMS integration",
        "Figma Plugins: Built productivity-enhancing Figma plugins to streamline design workflows",
        "AI Interview Simulator: Developed an AI-powered interview simulator using Google Gemini API",
        "Tournament Management System: Full-stack application with team management and payment integration (Paystack)"
    ]
    
    for proj in projects:
        proj_para = doc.add_paragraph(style='List Bullet')
        proj_run = proj_para.add_run(proj)
        proj_run.font.size = Pt(10)
        proj_run.font.name = 'Arial'
        proj_para.paragraph_format.left_indent = Inches(0.25)
        proj_para.space_after = Pt(2)
    
    # Save the document
    doc.save('c:/Users/HP/Documents/uxdimeji/OLADIMEJI_ABUBAKAR_RESUME.docx')
    print("Resume saved successfully as OLADIMEJI_ABUBAKAR_RESUME.docx")

if __name__ == "__main__":
    create_resume()
